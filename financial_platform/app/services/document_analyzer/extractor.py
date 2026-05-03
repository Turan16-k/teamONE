"""
ZIP'ten belge çıkarıcı.
PDF, Excel, CSV, Word, metin ve görsel dosyaları okur.
"""
from __future__ import annotations
import io
import csv
import zipfile
import logging
from pathlib import Path
from typing import List, Tuple, Optional

from .models import ExtractedDocument, DocumentType

logger = logging.getLogger(__name__)

_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"}
_EXT_MAP = {
    ".pdf":  DocumentType.PDF,
    ".xlsx": DocumentType.EXCEL,
    ".xls":  DocumentType.EXCEL,
    ".csv":  DocumentType.CSV,
    ".docx": DocumentType.WORD,
    ".doc":  DocumentType.WORD,
    ".txt":  DocumentType.TEXT,
    **{ext: DocumentType.IMAGE for ext in _IMAGE_EXTS},
}

_SKIP_PREFIXES = ("__MACOSX", ".", "_")


class ZipExtractor:
    MAX_FILE_BYTES = 50 * 1024 * 1024   # 50 MB / dosya
    MAX_FILES = 150

    def extract_all(self, zip_bytes: bytes) -> List[ExtractedDocument]:
        docs: List[ExtractedDocument] = []
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                entries = [
                    e for e in zf.infolist()
                    if not e.is_dir()
                    and not any(e.filename.startswith(p) for p in _SKIP_PREFIXES)
                    and Path(e.filename).suffix.lower() in _EXT_MAP
                ][: self.MAX_FILES]

                for entry in entries:
                    if entry.file_size > self.MAX_FILE_BYTES:
                        logger.warning("Büyük dosya atlandı: %s (%d MB)",
                                       entry.filename, entry.file_size // 1024 // 1024)
                        continue
                    try:
                        raw = zf.read(entry.filename)
                        doc = self._process(entry.filename, raw)
                        if doc:
                            docs.append(doc)
                    except Exception as exc:
                        logger.warning("İşlenemedi %s: %s", entry.filename, exc)

        except zipfile.BadZipFile as exc:
            raise ValueError(f"Geçersiz ZIP dosyası: {exc}")

        logger.info("%d belge çıkarıldı", len(docs))
        return docs

    # ── dosya tipi yönlendirme ─────────────────────────────────────────────

    def _process(self, filename: str, raw: bytes) -> Optional[ExtractedDocument]:
        suffix = Path(filename).suffix.lower()
        doc_type = _EXT_MAP.get(suffix, DocumentType.UNKNOWN)
        if doc_type == DocumentType.UNKNOWN:
            return None

        doc = ExtractedDocument(filename=Path(filename).name, file_type=doc_type)

        if doc_type == DocumentType.CSV:
            doc.text_content, doc.tables = self._csv(raw)
        elif doc_type == DocumentType.EXCEL:
            doc.text_content, doc.tables = self._excel(raw)
        elif doc_type == DocumentType.WORD:
            doc.text_content = self._word(raw)
        elif doc_type == DocumentType.TEXT:
            doc.text_content = self._text(raw)
        elif doc_type == DocumentType.PDF:
            doc.text_content = self._pdf(raw)
            if len(doc.text_content.strip()) < 80:   # büyük ihtimalle taranmış
                doc.raw_bytes = raw
        elif doc_type == DocumentType.IMAGE:
            doc.raw_bytes = raw

        return doc

    # ── format çıkarıcıları ───────────────────────────────────────────────

    def _csv(self, raw: bytes) -> Tuple[str, list]:
        try:
            text = raw.decode("utf-8-sig", errors="replace")
            rows = list(csv.reader(io.StringIO(text)))
            return "\n".join("\t".join(r) for r in rows[:200]), rows
        except Exception:
            return raw.decode("utf-8", errors="replace"), []

    def _excel(self, raw: bytes) -> Tuple[str, list]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts, all_rows = [], []
            for ws in wb.worksheets:
                rows = [
                    [str(v) if v is not None else "" for v in row]
                    for row in ws.iter_rows(max_row=300, values_only=True)
                    if any(v is not None for v in row)
                ]
                if rows:
                    all_rows.extend(rows)
                    parts.append(f"[Sayfa: {ws.title}]\n" + "\n".join("\t".join(r) for r in rows))
            return "\n\n".join(parts), all_rows
        except Exception as exc:
            logger.warning("Excel hatası: %s", exc)
            return f"[Excel okuma hatası: {exc}]", []

    def _word(self, raw: bytes) -> str:
        try:
            import docx as _docx
            doc = _docx.Document(io.BytesIO(raw))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            tbl_lines = []
            for tbl in doc.tables:
                for row in tbl.rows:
                    tbl_lines.append("\t".join(c.text for c in row.cells))
            return "\n".join(paras + tbl_lines)
        except Exception as exc:
            logger.warning("Word hatası: %s", exc)
            return f"[Word okuma hatası: {exc}]"

    def _text(self, raw: bytes) -> str:
        for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1254"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    def _pdf(self, raw: bytes) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            pages = []
            for page in reader.pages[:60]:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t)
            return "\n\n".join(pages)
        except Exception as exc:
            logger.warning("PDF metin çıkarma hatası: %s", exc)
            return ""
